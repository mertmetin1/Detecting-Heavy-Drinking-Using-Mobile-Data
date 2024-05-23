import sys
import pandas as pd
import itertools
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.ensemble import RandomForestRegressor
from sklearn.svm import SVR
from sklearn.metrics import mean_squared_error, mean_absolute_error
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io
import os
from PIL import Image

# Dosyaların bulunduğu dizin
directory = 'raw_tac'

# Dizin içindeki tüm dosyaları bir listeye ekle
file_list = [f for f in os.listdir(directory) if os.path.isfile(os.path.join(directory, f))]

for file in file_list:
    # Çıktıları ve grafikleri saklamak için bir StringIO nesnesi oluştur
    output = io.StringIO()

    # Standart çıktıyı ayarla
    sys.stdout = output

    # Analiz başlığı
    print("Analysis Of Person:", file)

    # Excel dosyasını oku
    tac_data = pd.read_excel(os.path.join(directory, file), sheet_name=0, skiprows=1)

    # Gerekli sütunları seç
    tac_data = tac_data[['TAC Level', 'IR Voltage', 'Temperature', 'Time', 'Date']]
    tac_data.dropna(inplace=True)

    # Veri setinin genel bilgilerini yazdır
    print("Veri Seti Başlıkları ve İlk 3 Satır:")
    print(tac_data.tail(3))
    print("\nVeri Seti Bilgileri:")
    print(tac_data.info())
    print("\nVeri Seti İstatistikleri:")
    print(tac_data.describe())
    print("\nEksik Değerlerin Sayısı:")
    print(tac_data.isnull().sum())
    print("\nBenzersiz Değerlerin Sayısı:")
    print(tac_data.nunique())

    # Özellikleri ve hedef değişkenleri ayır
    y = tac_data['TAC Level']
    X = tac_data[['IR Voltage', 'Temperature']]

    # Min-max normalizasyon uygula
    scaler = MinMaxScaler()
    X_scaled = pd.DataFrame(scaler.fit_transform(X), columns=X.columns)

    # Korelasyon matrisi
    X_scaled_cm = X_scaled.corr()
    print("\nCorrelation Matrix:")
    print(X_scaled_cm)

    # Korelasyon matrisini görselleştir
    plt.figure(figsize=(10, 8))
    sns.heatmap(X_scaled_cm, annot=True, cmap='coolwarm', fmt=".2f")
    plt.title('Özelliklerin Korelasyon Matrisi')
    plt.xlabel('Özellikler')
    plt.ylabel('Özellikler')
    correlation_matrix_file = file+"_correlation_matrix.png"
    plt.savefig(correlation_matrix_file)
    plt.close()

    # Eğitim ve test setlerine veriyi ayır
    X_train, X_test, y_train, y_test = train_test_split(X_scaled, y, test_size=0.3, random_state=42)

    # Lineer Regresyon
    print("###################################################### Lineer Regresyon #############################################")

    lr_model = LinearRegression()
    lr_model.fit(X_train, y_train)
    lr_y_pred = lr_model.predict(X_test)

    lr_mse = mean_squared_error(y_test, lr_y_pred)
    lr_accuracy = lr_model.score(X_test, y_test)
    lr_mae = mean_absolute_error(y_test, lr_y_pred)
    lr_rmse = mean_squared_error(y_test, lr_y_pred, squared=False)

    print("Lineer Regresyon Ortalama Kare Hatası:", lr_mse)
    print("Lineer Regresyon Doğruluk:", lr_accuracy)
    print("Lineer Regresyon Ortalama Mutlak Hata:", lr_mae)
    print("Lineer Regresyon Kök Ortalama Kare Hata:", lr_rmse)

    print("Lineer Regresyon - Test ve Tahmin Değerleri:")
    print("Test Değerleri:", y_test.values)
    print("Tahmin Edilen Değerler:", lr_y_pred)

    plt.figure(figsize=(12, 6))
    plt.scatter(y_test, lr_y_pred, label="LineerRegresyon Tahminleri")
    plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], color='red', linestyle='--', label='Regresyon Çizgisi')
    plt.xlabel('Gerçek TAC')
    plt.ylabel('Tahmin Edilen TAC')
    plt.title('LineerRegresyon Tahminleri ve Gerçek Değerler')
    plt.legend()
    linear_regression_plot_file = file+"_linear_regression_plot.png"
    plt.savefig(linear_regression_plot_file)
    plt.close()

    # Random Forest Regresyonu
    print("###################################################### Random Forest Regresyonu #############################################")

    rf_param_grid = {
        'n_estimators': [50, 100, 200],
        'max_depth': [None, 10, 20],
        'min_samples_split': [2, 5, 10],
        'min_samples_leaf': [1, 2, 4]
    }

    best_rf_mse = float('inf')
    best_rf_model = None

    for rf_params in itertools.product(*rf_param_grid.values()):
        rf_regressor = RandomForestRegressor(n_estimators=rf_params[0], max_depth=rf_params[1],
                                            min_samples_split=rf_params[2], min_samples_leaf=rf_params[3], random_state=42)
        rf_regressor.fit(X_train, y_train)
        rf_y_pred = rf_regressor.predict(X_test)
        rf_mse = mean_squared_error(y_test, rf_y_pred)
        
        if rf_mse < best_rf_mse:
            best_rf_mse = rf_mse
            best_rf_model = rf_regressor
        
            print("RandomForestRegressor Parameters:", rf_params)
            print("Mean Squared Error:", rf_mse)

    rf_accuracy = best_rf_model.score(X_test, y_test)
    rf_r_squared = best_rf_model.score(X_test, y_test)
    rf_mae = mean_absolute_error(y_test, best_rf_model.predict(X_test))
    rf_rmse = mean_squared_error(y_test, best_rf_model.predict(X_test), squared=False)

    print("RandomForestRegressor Accuracy:", rf_accuracy)
    print("RandomForestRegressor R-squared:", rf_r_squared)
    print("RandomForestRegressor Mean Absolute Error:", rf_mae)
    print("RandomForestRegressor Root Mean Squared Error:", rf_rmse)
    print("Random Forest Regressor - Test and Prediction Values:")
    print("Test Values: ", y_test.values)
    print("Predicted Values:", best_rf_model.predict(X_test))

    plt.figure(figsize=(12, 6))
    plt.scatter(y_test, best_rf_model.predict(X_test), label="RandomForestRegressor Predictions")
    plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], color='red', linestyle='--', label='Regression Line')
    plt.xlabel('Actual TAC')
    plt.ylabel('Predicted TAC')
    plt.title('RandomForestRegressor Predictions vs Actual Values')
    plt.legend()
    random_forest_regression_plot_file = file+"_random_forest_regression_plot.png"
    plt.savefig(random_forest_regression_plot_file)
    plt.close()

    # SVR
    print("###################################################### SVR #############################################")

    param_grid = {
        'kernel': ['linear', 'poly', 'rbf', 'sigmoid'],
        'C': [0.1, 1, 10],
        'gamma': ['scale', 'auto'],
        'epsilon': [0.1, 0.01, 0.001]
    }

    best_mse = float('inf')
    best_model = None

    for params in itertools.product(*param_grid.values()):
        svr = SVR(kernel=params[0], C=params[1], gamma=params[2], epsilon=params[3])
        svr.fit(X_train, y_train)
        svr_y_pred = svr.predict(X_test)
        mse = mean_squared_error(y_test, svr_y_pred)
        
        if mse < best_mse:
            best_mse = mse
            best_model = svr
        
            print("Parameters:", params)
            print("Mean Squared Error:", mse)

    svr_accuracy = best_model.score(X_test, y_test)
    svr_r_squared = best_model.score(X_test, y_test)
    svr_mae = mean_absolute_error(y_test, best_model.predict(X_test))
    svr_rmse = mean_squared_error(y_test, best_model.predict(X_test), squared=False)

    print("SVR Accuracy:", svr_accuracy)
    print("SVR R-squared:", svr_r_squared)
    print("SVR Mean Absolute Error:", svr_mae)
    print("SVR Root Mean Squared Error:", svr_rmse)

    print("SVR - Test and Prediction Values:")
    print("Test Values:", y_test.values)
    print("Predicted Values:", best_model.predict(X_test))

    plt.figure(figsize=(12, 6))
    plt.scatter(y_test, svr_y_pred, label="SVR Predictions")
    plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], color='red', linestyle='--', label='Regression Line')
    plt.xlabel('Actual TAC')
    plt.ylabel('Predicted TAC')
    plt.title('SVR Predictions vs Actual Values')
    plt.legend()
    svr_plot_file = file+"_svr_plot.png"
    plt.savefig(svr_plot_file)
    plt.close()

    # Classification thresholds
    legal_limit = 0.08
    mean_tac = 0.065
    max_tac = 0.443
    inner_quartiles = [0.002, 0.029, 0.092]

    # TAC Classification
    def classify_tac(tac_level):
        if tac_level < inner_quartiles[0]:
            return "Less Than Legal Limit"
        elif inner_quartiles[0] <= tac_level <= inner_quartiles[2]:
            return "About Legal Limit"
        else:
            return "Illegal: Heavy Alcohol"

    lr_classified_tac = [classify_tac(tac_level) for tac_level in lr_y_pred]
    rf_classified_tac = [classify_tac(tac_level) for tac_level in best_rf_model.predict(X_test)]
    svr_classified_tac = [classify_tac(tac_level) for tac_level in best_model.predict(X_test)]

    results = pd.DataFrame({
        'Time': tac_data.iloc[X_test.index]['Time'],
        'Actual_TAC': y_test,
        'LR_Predicted_TAC': lr_y_pred,
        'RF_Predicted_TAC': best_rf_model.predict(X_test),
        'SVR_Predicted_TAC': best_model.predict(X_test),
        'LR_Classified_TAC': lr_classified_tac,
        'RF_Classified_TAC': rf_classified_tac,
        'SVR_Classified_TAC': svr_classified_tac
    })

    results.sort_values(by='Time', inplace=True)

    # Sonuçları yazdır
    print("\nResults DataFrame for Person", file, ":\n", results)

    # DOCX dosyası oluştur
    doc = Document()
    doc.add_heading('Alcohol Consumption Analysis Report', 0)

    # Ekleme işlemleri
    doc.add_heading(file+" Alcohol Consumption Analysis Report", level=1)
    doc.add_heading('Data Summary', level=1)
    doc.add_paragraph("Veri Seti Başlıkları ve İlk 3 Satır:\n")
    doc.add_paragraph(str(tac_data.tail(3)))
    doc.add_paragraph("\nVeri Seti Bilgileri:\n")
    doc.add_paragraph(str(tac_data.info()))
    doc.add_paragraph("\nVeri Seti İstatistikleri:\n")
    doc.add_paragraph(str(tac_data.describe()))
    doc.add_paragraph("\nEksik Değerlerin Sayısı:\n")
    doc.add_paragraph(str(tac_data.isnull().sum()))
    doc.add_paragraph("\nBenzersiz Değerlerin Sayısı:\n")
    doc.add_paragraph(str(tac_data.nunique()))

    doc.add_heading('Correlation Matrix of Features', level=1)
    doc.add_picture(correlation_matrix_file, width=Inches(6))

    doc.add_heading('Linear Regression Results', level=1)
    doc.add_paragraph("Mean Squared Error: " + str(lr_mse))
    doc.add_paragraph("Accuracy: " + str(lr_accuracy))
    doc.add_paragraph("Mean Absolute Error: " + str(lr_mae))
    doc.add_paragraph("Root Mean Squared Error: " + str(lr_rmse))
    doc.add_paragraph("\nLinear Regression - Test and Prediction Values:\n")
    doc.add_paragraph("Test Values: " + str(y_test.values))
    doc.add_paragraph("Predicted Values: " + str(lr_y_pred))
    doc.add_picture(linear_regression_plot_file, width=Inches(6))

    doc.add_heading('Random Forest Regression Results', level=1)
    doc.add_paragraph("Accuracy: " + str(rf_accuracy))
    doc.add_paragraph("R-squared: " + str(rf_r_squared))
    doc.add_paragraph("Mean Absolute Error: " + str(rf_mae))
    doc.add_paragraph("Root Mean Squared Error: " + str(rf_rmse))
    doc.add_paragraph("\nRandom Forest Regressor - Test and Prediction Values:\n")
    doc.add_paragraph("Test Values: " + str(y_test.values))
    doc.add_paragraph("Predicted Values: " + str(best_rf_model.predict(X_test)))
    doc.add_picture(random_forest_regression_plot_file, width=Inches(6))

    doc.add_heading('SVR Results', level=1)
    doc.add_paragraph("Accuracy: " + str(svr_accuracy))
    doc.add_paragraph("R-squared: " + str(svr_r_squared))
    doc.add_paragraph("Mean Absolute Error: " + str(svr_mae))
    doc.add_paragraph("Root Mean Squared Error: " + str(svr_rmse))
    doc.add_paragraph("\nSVR - Test and Prediction Values:\n")
    doc.add_paragraph("Test Values: " + str(y_test.values))
    doc.add_paragraph("Predicted Values: " + str(best_model.predict(X_test)))
    doc.add_picture(svr_plot_file, width=Inches(6))

    doc.add_heading('Results DataFrame', level=1)
    doc.add_paragraph("Results DataFrame for Person " + file + ":\n")
    results_table = doc.add_table(rows=results.shape[0]+1, cols=results.shape[1])
    for j in range(results.shape[-1]):
        results_table.cell(0, j).text = results.columns[j]
    for i in range(results.shape[0]):
        for j in range(results.shape[-1]):
            results_table.cell(i+1, j).text = str(results.values[i,j])

    doc.add_page_break()

    # DOCX dosyasını kaydet
    doc.save(file+"_Alcohol_Consumption_Analysis_Report.docx")

    # Standart çıktıyı geri yükle
    sys.stdout = sys.__stdout__

    # StringIO nesnesini okuma konumunu başlangıca getir
    output.seek(0)

    # Çıktıları ekrana yazdır
    print(output.read())

    # PNG dosyalarını sil
    os.remove(correlation_matrix_file)
    os.remove(linear_regression_plot_file)
    os.remove(random_forest_regression_plot_file)
    os.remove(svr_plot_file)



